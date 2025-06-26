from docx import Document as DocxDocument
from langchain.schema import Document
import os

def load_docx(file_path, source_name=None):
    docx = DocxDocument(file_path)
    text = "\n".join([p.text for p in docx.paragraphs])
    # Optionally convert to Markdown here (stub)
    metadata = {"source": source_name or os.path.basename(file_path)}
    return [Document(page_content=text, metadata=metadata)] 